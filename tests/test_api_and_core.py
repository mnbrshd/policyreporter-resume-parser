import io
import tempfile
import os
import asyncio
import pytest

from fastapi import UploadFile
from fastapi.testclient import TestClient
from pathlib import Path

# Import your FastAPI app
from resume_parser.app import app

# Import internal classes to test parsing logic
from resume_parser.parsers.pdf_parser import PDFParser
from resume_parser.parsers.word_parser import WordParser
from resume_parser.extractors.name_extractor import NameExtractor
from resume_parser.extractors.email_extractor import EmailExtractor
from resume_parser.extractors.skills_extractor import SkillsExtractor
from resume_parser.extractors.llm_client.clients import LLMClient, OpenAILLMClient
from resume_parser.resume_extractor import ResumeExtractor
from resume_parser.framework import ResumeParserFramework
from resume_parser.models import ResumeData


@pytest.fixture
def core_framework() -> ResumeParserFramework:
    llm_client = OpenAILLMClient()
    extractors = {
        "name": NameExtractor(),
        "email": EmailExtractor(),
        "skills": SkillsExtractor(llm_client=llm_client),
    }
    parsers = {
        ".pdf": PDFParser(),
        ".docx": WordParser(),
    }
    return ResumeParserFramework(parsers=parsers, extractor=ResumeExtractor(extractors))


def test_core_framework(core_framework: ResumeParserFramework):
    fake_file = "tests/test_resume.pdf"

    resume: ResumeData = asyncio.run(core_framework.parse_resume(fake_file))
    data = resume.model_dump()

    assert data["name"] == "Muhammad Muneeb Arshad"
    assert data["email"] == "mnbrshd@gmail.com"
    assert data["skills"], "skills list must not be empty"
    assert isinstance(data.get("skills"), list), "skills must be a list"


def test_core_framework_unsupported_suffix(tmp_path, core_framework: ResumeParserFramework):
    fake_file = tmp_path / "resume.txt"
    fake_file.write_text("Some content", encoding="utf-8")
    with pytest.raises(ValueError) as excinfo:
        core_framework.parse_resume(str(fake_file))
    assert "Unsupported file suffix" in str(excinfo.value)


@pytest.fixture
def client() -> TestClient:
    return TestClient(app)


def test_parse_resume_endpoint_missing_file(client: TestClient):
    response = client.post("/parse-resume/", data={})
    assert response.status_code == 422


def test_parse_resume_endpoint_invalid_file_type(client: TestClient):
    # Upload a file with unsupported extension
    files = {"file": ("resume.txt", io.BytesIO(b"Dummy resume"), "text/plain")}
    response = client.post("/parse-resume/", files=files)
    assert response.status_code == 400
    assert "Unsupported file type" in response.json().get("detail", "")


def test_parse_resume_endpoint_success_docx(client: TestClient, tmp_path):
    # Create a temporary .docx file for testing
    from docx import Document

    doc = Document()
    doc.add_paragraph("Bob Example")
    doc.add_paragraph("Contact: bob.example@example.com")
    doc.add_paragraph("Skills: Python, Docker, Kubernetes")
    tmp_file_path = tmp_path / "resume.docx"
    doc.save(str(tmp_file_path))

    with open(tmp_file_path, "rb") as f:
        files = {"file": ("resume.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        response = client.post("/parse-resume/", files=files)

    assert response.status_code == 200
    json_data = response.json()
    assert json_data["name"] == "Bob Example"
    assert json_data["email"] == "bob.example@example.com"
    assert "Python" in json_data["skills"]
    assert "Docker" in json_data["skills"]


def test_parse_resume_endpoint_success_pdf(client: TestClient, tmp_path):
    fake_file = "tests/test_resume.pdf"

    with open(fake_file, "rb") as f:
        files = {"file": ("resume.pdf", f, "application/pdf")}
        response = client.post("/parse-resume/", files=files)

    assert response.status_code == 200
    json_data = response.json()
    assert json_data.get("name") in ("Muhammad Muneeb Arshad", "")
    assert "email" in json_data and "skills" in json_data
